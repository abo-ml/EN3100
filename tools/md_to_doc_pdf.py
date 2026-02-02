#!/usr/bin/env python3
"""Convert Markdown to Word (.docx) and optionally PDF.

Usage:
    python tools/md_to_doc_pdf.py --md reports/dissertation.md \
        --docx reports/EN3100_Dissertation.docx \
        --pdf reports/EN3100_Dissertation.pdf

Note: Math expressions stay as raw LaTeX in the Word output.
For camera-ready PDF with rendered math, use Pandoc with LaTeX.
"""
from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def markdown_to_docx(md_path: Path, docx_path: Path) -> None:
    """Convert Markdown file to Word document with simple Calibri 11 styling."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        line = lines[i].rstrip("\n")

        # Skip empty lines
        if not line.strip():
            doc.add_paragraph("")
            i += 1
            continue

        # Headings
        if line.startswith("# "):
            p = doc.add_heading(line[2:], level=1)
            i += 1
            continue
        if line.startswith("## "):
            p = doc.add_heading(line[3:], level=2)
            i += 1
            continue
        if line.startswith("### "):
            p = doc.add_heading(line[4:], level=3)
            i += 1
            continue
        if line.startswith("#### "):
            p = doc.add_heading(line[5:], level=4)
            i += 1
            continue

        # Bullet lists
        if line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(line[2:], style="List Bullet")
            i += 1
            continue

        # Numbered lists
        match = re.match(r"^(\d+)\.\s+(.*)$", line)
        if match:
            p = doc.add_paragraph(match.group(2), style="List Number")
            i += 1
            continue

        # Code blocks
        if line.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i].rstrip("\n"))
                i += 1
            i += 1  # skip closing ```
            p = doc.add_paragraph()
            run = p.add_run("\n".join(code_lines))
            run.font.name = "Courier New"
            run.font.size = Pt(10)
            continue

        # Block quotes
        if line.startswith("> "):
            p = doc.add_paragraph(line[2:])
            p.paragraph_format.left_indent = Pt(36)
            p.italic = True
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph(line)
        i += 1

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))
    print(f"Saved Word document to {docx_path}")


def docx_to_pdf_fallback(docx_path: Path, pdf_path: Path) -> None:
    """Create a simple text-based PDF fallback.

    This is a basic fallback that writes text to PDF.
    For proper formatting, use Pandoc with LaTeX.
    """
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        from reportlab.lib.units import inch

        doc = Document(str(docx_path))
        c = canvas.Canvas(str(pdf_path), pagesize=letter)
        width, height = letter
        y = height - inch

        for para in doc.paragraphs:
            text = para.text
            if not text.strip():
                y -= 12
                continue

            # Simple text wrapping
            words = text.split()
            line = ""
            for word in words:
                test_line = f"{line} {word}".strip()
                if len(test_line) > 80:
                    c.drawString(inch, y, line)
                    y -= 14
                    line = word
                    if y < inch:
                        c.showPage()
                        y = height - inch
                else:
                    line = test_line

            if line:
                c.drawString(inch, y, line)
                y -= 14

            if y < inch:
                c.showPage()
                y = height - inch

        c.save()
        print(f"Saved PDF to {pdf_path}")

    except ImportError:
        print("reportlab not installed. Skipping PDF generation.")
        print("For better PDF output, use: pandoc input.md -o output.pdf")


def main() -> None:
    parser = argparse.ArgumentParser(description="Convert Markdown to Word and optionally PDF")
    parser.add_argument("--md", required=True, help="Input Markdown file")
    parser.add_argument("--docx", required=True, help="Output Word document path")
    parser.add_argument("--pdf", default=None, help="Optional PDF output path")
    args = parser.parse_args()

    md_path = Path(args.md)
    docx_path = Path(args.docx)

    if not md_path.exists():
        print(f"Error: Markdown file not found: {md_path}")
        return

    markdown_to_docx(md_path, docx_path)

    if args.pdf:
        pdf_path = Path(args.pdf)
        docx_to_pdf_fallback(docx_path, pdf_path)


if __name__ == "__main__":
    main()
